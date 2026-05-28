import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_astro_document():
    # 1. Initialize Document
    doc = Document()
    
    # 2. Configure Global Page Margins (Standard 1 Inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 3. Helper Functions for Custom Styling
    def add_styled_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        return heading

    def add_code_block(code_text):
        # Creates a specialized paragraph with background shading and a monospace font
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        # Apply light gray background tint (XML Shading)
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="F4F4F6"/>'
        p._p.get_or_add_pPr().append(parse_xml(shading_xml))
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E) # Deep slate text
        return p

    # --- DOCUMENT GENERATION ---

    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)
    run_title = title.add_run("Documentation: Architecting a Content-Driven Blog Using Astro")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run_sub = subtitle.add_run("A Reference Guide for DevOps & Cloud Engineers")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # Section 1
    add_styled_heading("1. Executive Summary & Core Philosophy", level=1, space_before=18)
    p1 = doc.add_paragraph(
        "Traditional Content Management Systems introduce heavy database overhead, complex caching layers, "
        "and expanded security attack surfaces. For a DevOps Engineer, the core architectural goal is to treat "
        "Content as Code. Astro achieves this by generating pure, static HTML at build time (Zero-JS by default), "
        "shifting the dynamic rendering workload entirely from the runtime environment to the compilation phase."
    )
    p1.paragraph_format.space_after = Pt(12)

    # Section 2
    add_styled_heading("2. Infrastructure & Directory Blueprint", level=1)
    p2 = doc.add_paragraph(
        "A production-ready Astro repository maintains a clean separation between structural layouts, data validation schemas, "
        "and pure Markdown source contents. Below is the standard structural framework layout used for development:"
    )
    p2.paragraph_format.space_after = Pt(6)

    directory_tree = (
        "├── .github/workflows/      # CI/CD Automation runner files\n"
        "├── public/\n"
        "│   ├── admin/\n"
        "│   │   └── config.yml      # Decap CMS Administrative layout structure\n"
        "│   └── townscribe-img.jpg  # Static global assets\n"
        "├── src/\n"
        "│   ├── components/         # Reusable structural layout blocks\n"
        "│   ├── content/\n"
        "│   │   ├── config.ts       # Type Safety schemas for Markdown validation\n"
        "│   │   └── blog/           # Pure Markdown (.md) posts written by authors\n"
        "│   ├── layouts/            # Global HTML page shells\n"
        "│   └── pages/              # File-based routing layout mapping\n"
        "│       ├── index.astro     # Automated homepage grid loop\n"
        "│       └── blog/\n"
        "│           └── [...slug].astro\n"
        "├── astro.config.mjs        # Core compilation configurations\n"
        "└── package.json            # Node runtime dependencies"
    )
    add_code_block(directory_tree)

    # Section 3
    add_styled_heading("3. Implementation Blueprint (Step-by-Step)", level=1)
    
    add_styled_heading("Step 3.1: Defining the Content Schema (Type Safety)", level=2)
    doc.add_paragraph(
        "To prevent formatting variances from failing compilation builds, leverage Astro’s native Zod validation engine "
        "inside the src/content/config.ts configuration directory:"
    )
    
    schema_code = (
        "import { defineCollection, z } from 'astro:content';\n\n"
        "const blog = defineCollection({\n"
        "    type: 'content',\n"
        "    schema: z.object({\n"
        "        title: z.string(),\n"
        "        description: z.string(),\n"
        "        pubDate: z.coerce.date(),\n"
        "        heroImage: z.string().optional(),\n"
        "        category: z.enum(['news', 'politics', 'business', 'sport']),\n"
        "        seoTitle: z.string().optional(),\n"
        "        seoDescription: z.string().optional(),\n"
        "    }),\n"
        "});\n\n"
        "export const collections = { blog };"
    )
    add_code_block(schema_code)

    add_styled_heading("Step 3.2: Setting Up the Dynamic Query Loop", level=2)
    doc.add_paragraph(
        "In your index pages, execute an asynchronous data fetch at compilation time to collect, clean, "
        "and order your content elements sequentially:"
    )

    query_code = (
        "---\n"
        "import { getCollection } from 'astro:content';\n"
        "import BaseHead from '../components/BaseHead.astro';\n\n"
        "// Async data fetch triggered strictly at build time\n"
        "const posts = (await getCollection('blog')).sort(\n"
        "    (a, b) => b.data.pubDate.valueOf() - a.data.pubDate.valueOf()\n"
        ");\n"
        "---\n"
        "<ul class=\"blog-grid\">\n"
        "    {posts.map((post) => (\n"
        "        <li>\n"
        "            <a href={`/blog/${post.slug}/`}>\n"
        "                <img src={post.data.heroImage} alt=\"\" />\n"
        "                <h4>{post.data.title}</h4>\n"
        "            </a>\n"
        "        </li>\n"
        "    ))}\n"
        "</ul>"
    )
    add_code_block(query_code)

    # Section 4
    add_styled_heading("4. Decoupled Identity & Decap CMS Integration", level=1)
    doc.add_paragraph(
        "Write operations are handled securely without a persistent live database cluster by leveraging "
        "Decap CMS paired with a DecapBridge gateway. Put the configuration settings directly inside public/admin/config.yml:"
    )

    decap_config = (
        "backend:\n"
        "  name: github\n"
        "  repo: your-username/your-repo-name\n"
        "  branch: main\n"
        "  base_url: https://your-decap-bridge-instance.comn"
        "  auth_endpoint: /auth\n\n"
        "local_backend: true\n"
        "media_folder: \"public/uploads\"\n"
        "public_folder: \"/uploads\"\n\n"
        "collections:\n"
        "  - name: \"blog\"\n"
        "    label: \"Articles Feed\"\n"
        "    folder: \"src/content/blog\"\n"
        "    create: true\n"
        "    fields:\n"
        "      - { label: \"Title\", name: \"title\", widget: \"string\" }\n"
        "      - { label: \"Publish Date\", name: \"pubDate\", widget: \"datetime\" }\n"
        "      - { label: \"Category\", name: \"category\", widget: \"select\", options: [\"news\", \"politics\", \"business\", \"sport\"] }\n"
        "      - { label: \"Body Content\", name: \"body\", widget: \"markdown\" }"
    )
    add_code_block(decap_config)

    # Section 5
    add_styled_heading("5. Deployment Blueprint: Cloudflare Pages Edge Pipeline", level=1)
    doc.add_paragraph(
        "Deploying across an Edge Cloud Network (Cloudflare Pages) optimizes global TTFB. Define an autonomous runner file "
        "under .github/workflows/deploy.yml to automate build tasks upon commit hooks:"
    )

    actions_code = (
        "name: Production Edge Deployment Pipeline\n\n"
        "on:\n"
        "  push:\n"
        "    branches: [ main ]\n\n"
        "jobs:\n"
        "  deploy:\n"
        "    runs-on: ubuntu-latest\n"
        "    steps:\n"
        "      - uses: actions/checkout@v4\n"
        "      - uses: actions/setup-node@v4\n"
        "        with:\n"
        "          node-version: 20\n"
        "          cache: 'npm'\n"
        "      - run: npm ci\n"
        "      - run: npm run build\n"
        "      - name: Publish Direct to Cloudflare Pages\n"
        "        uses: cloudflare/pages-action@v1\n"
        "        with:\n"
        "          apiToken: ${{ secrets.CLOUDFLARE_API_TOKEN }}\n"
        "          accountId: ${{ secrets.CLOUDFLARE_ACCOUNT_ID }}\n"
        "          projectName: 'townscribe-news'\n"
        "          directory: './dist'"
    )
    add_code_block(actions_code)

    # Section 6
    add_styled_heading("6. Optimization Metrics for Operations (DevOps Checklist)", level=1)
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r1 = bp1.add_run("Cache Invalidation: ")
    r1.bold = True
    bp1.add_run("Enforce aggressive immutable asset caching profiles; edge platforms handle instant cache drops automatically on each green-lit deployment execution loop.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    r2 = bp2.add_run("Identity Separation: ")
    r2.bold = True
    bp2.add_run("Isolate content writers from accessing infrastructure directly. Utilize isolated, OAuth middleware gates (DecapBridge classic auth) to restrict access limits to standard markdown paths only.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    r3 = bp3.add_run("Performance Metrics: ")
    r3.bold = True
    bp3.add_run("Leverage Astro's build-time dynamic picture optimization structures to force automatic asset transformations down to next-generation formats (.webp), maintaining near perfect Lighthouse optimization thresholds.")

    # Save document
    filename = "Astro_DevOps_Blog_Architecture.docx"
    doc.save(filename)
    print(f"Success! Technical guide compiled cleanly into file: {filename}")

if __name__ == "__main__":
    create_astro_document()